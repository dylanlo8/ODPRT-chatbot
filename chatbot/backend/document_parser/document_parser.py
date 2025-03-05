import cv2
import docx
import io
import json
import mimetypes
import numpy as np
import os
import pdfplumber
import requests
import shutil
from docx import Document
from docx2pdf import convert
from fpdf import FPDF
import pickle
import pytesseract
from langchain_experimental.text_splitter import SemanticChunker
from pdf2image import convert_from_path
from PIL import Image
from typing import List
from chatbot.backend.services.logger import logger
from chatbot.backend.services.models.embedding_model import embedding_model
from chatbot.backend.services.models.models import vlm


class DocumentParser:
    """class to parse attachments and return ingestable chunks of information"""

    def __init__(
        self,
        attachment_directory: str = "processed_docs/emails_with_attachments",
        similarity_threshold: float = 0.85,
    ):
        self.directory = attachment_directory
        self.similarity_threshold = similarity_threshold
        self.embedding_model = embedding_model
        self.logger = logger

    def read_attachments(self):
        """Reads all attachments in the directory and returns a list of file paths."""
        files = []
        for root, _, filenames in os.walk(self.directory):
            for filename in filenames:
                files.append(os.path.join(root, filename))
        return files

    def filter_useful_attachments(self) -> List[str]:
        """
        Filters and processes useful attachments from emails in 'processed_docs/emails_with_attachments'.

        Returns:
            List[str]: List of relevant processed attachment file paths.
        """
        useful_attachments = []

        if not os.path.exists(self.directory):
            self.logger.warning(f"Directory {self.directory} does not exist.")
            return []

        for email_folder in os.listdir(self.directory):
            email_path = os.path.join(self.directory, email_folder)
            original_attachments_path = os.path.join(email_path, "original_attachments")
            processed_attachments_path = os.path.join(
                email_path, "processed_attachments"
            )
            cleaned_email_path = os.path.join(email_path, "cleaned_msg.txt")

            if not os.path.isdir(email_path):
                continue

            if not os.path.exists(cleaned_email_path):
                self.logger.warning(f"Missing cleaned email: {cleaned_email_path}")
                continue

            with open(cleaned_email_path, "r", encoding="utf-8") as f:
                email_thread = f.read().strip()

            os.makedirs(processed_attachments_path, exist_ok=True)

            for file in os.listdir(original_attachments_path):
                attachment_path = os.path.join(original_attachments_path, file)
                file_ext = file.lower().split(".")[-1]

                if file == "cleaned_msg.txt":
                    continue

                processed_file_path = ""

                if file_ext in ["png", "jpg", "jpeg", "bmp", "tiff"]:
                    payload = vlm._build_payloads_for_attachments(
                        email_thread, [attachment_path]
                    )[0]
                    response = requests.post(vlm.api, headers=vlm.headers, json=payload)

                    try:
                        response_json = response.json()
                        answer = (
                            response_json.get("choices", [{}])[0]
                            .get("message", {})
                            .get("content", "")
                        )

                        if answer.strip().lower() == "relevant":
                            processed_file_path = os.path.join(
                                processed_attachments_path, file
                            )
                            shutil.copy(attachment_path, processed_file_path)

                    except json.JSONDecodeError:
                        self.logger.error(
                            f"Invalid JSON response for {attachment_path}: {response.text}"
                        )

                elif file_ext == "docx":
                    # Convert DOCX to PDF first
                    pdf_path = self.convert_docx_to_pdf(attachment_path)
                    relevant_images = self.classify_attachment_relevance(
                        email_thread, pdf_path
                    )

                    # Save only relevant pages
                    if relevant_images:
                        processed_file_path = os.path.join(processed_attachments_path, os.path.basename(relevant_images))
                        # shutil.copy(relevant_images, processed_file_path)

                elif file_ext == "pdf":
                    # Classify pages directly from the PDF
                    relevant_images = self.classify_attachment_relevance(
                        email_thread, attachment_path
                    )

                    # Save only relevant pages
                    if relevant_images:
                        processed_file_path = os.path.join(processed_attachments_path, os.path.basename(relevant_images))
                        # shutil.copy(relevant_images, processed_file_path)


                # If a processed file was created, add it to the list of useful attachments
                if processed_file_path:
                    useful_attachments.append(processed_file_path)

                # Cleanup intermediate processing files (like extracted images)
                temp_image_dirs = [os.path.join(email_path, "temp_images")]
                for temp_dir in temp_image_dirs:
                    shutil.rmtree(temp_dir, ignore_errors=True)

        self.logger.info(f"Total useful attachments: {len(useful_attachments)}")
        return useful_attachments

    def convert_docx_to_pdf(self, docx_path: str) -> str:
        """
        Converts a DOCX file to PDF and saves it in a temporary directory.

        Args:
            docx_path (str): Path to the DOCX file.

        Returns:
            str: Path to the generated temporary PDF file.
        """
        temp_pdf_dir = os.path.join(os.path.dirname(os.path.dirname(docx_path)), "temp_pdfs")  # Store temporary PDFs separately
        os.makedirs(temp_pdf_dir, exist_ok=True)  # Ensure directory exists

        output_pdf_path = os.path.join(temp_pdf_dir, os.path.basename(docx_path).replace(".docx", ".pdf"))

        print(f"Converting DOCX to PDF: {output_pdf_path}")
        convert(docx_path, output_pdf_path)

        return output_pdf_path

    def convert_pdf_to_images(self, pdf_path: str, output_folder: str = None, dpi: int = 300) -> List[str]:
        """
        Converts a PDF into images, one image per page.

        Args:
            pdf_path (str): Path to the PDF file.
            output_folder (str, optional): Folder to save extracted images.
            dpi (int, optional): Resolution of the output images.

        Returns:
            List[str]: List of paths to the saved image files.
        """
        if output_folder is None:
            output_folder = os.path.join(os.path.dirname(os.path.dirname(pdf_path)), "temp_images")
        os.makedirs(output_folder, exist_ok=True)
        poppler_path = poppler_path = "/opt/homebrew/bin"
        images = convert_from_path(pdf_path, dpi=dpi, poppler_path = poppler_path)
        image_paths = []

        for i, image in enumerate(images):
            image_path = os.path.join(output_folder, f"page_{i+1}.png")
            image.save(image_path, "PNG")
            image_paths.append(image_path)

        return image_paths

    def classify_attachment_relevance(
        self, email_thread: str, attachment_path: str
    ) -> str:
        """
        Classifies whether an attachment (PDF, DOCX, or images) is relevant based on the email context.

        Args:
            email_thread (str): The cleaned email text.
            attachment_path (str): The file path to the attachment.

        Returns:
            str: Path to the final processed relevant PDF, or "not_relevant" if no relevant content is found.
        """
        file_ext = attachment_path.lower().split(".")[-1]
        useful_images = []
        temp_pdf_path = None  # Track temporary PDF

        email_dir = os.path.dirname(os.path.dirname(attachment_path))  # Get email folder
        processed_attachments_path = os.path.join(email_dir, "processed_attachments")  # Store processed PDFs here
        os.makedirs(processed_attachments_path, exist_ok=True)  # Ensure directory exists

        # Convert DOCX to PDF first, then classify pages
        if file_ext in ["docx"]:
            temp_pdf_path = self.convert_docx_to_pdf(attachment_path)  # Convert DOCX → PDF
            if not temp_pdf_path:
                return "not_relevant"
            images = self.convert_pdf_to_images(temp_pdf_path)  # Convert PDF → Images
        elif file_ext in ["pdf"]:
            images = self.convert_pdf_to_images(attachment_path)  # Convert PDF → Images
        elif file_ext in ["png", "jpg", "jpeg", "bmp", "tiff"]:
            images = [attachment_path]  # It's already an image
        else:
            self.logger.warning(f"Unsupported file format: {attachment_path}")
            return "not_relevant"

        payloads = vlm._build_payloads_for_attachments(email_thread, images)

        for image_path, payload in zip(images, payloads):
            response = requests.post(vlm.api, headers=vlm.headers, json=payload)

            try:
                response_json = response.json()
                answer = (
                    response_json.get("choices", [{}])[0]
                    .get("message", {})
                    .get("content", "")
                )

                if not answer.strip():
                    self.logger.error(f"Empty response for {image_path}. Skipping.")
                    continue

                self.logger.info(f"Raw model response for {image_path}: {answer}")

                # Parse JSON response
                try:
                    json_answer = json.loads(answer)
                    classification = json_answer.get("classification", "not_relevant")

                    if classification == "relevant":
                        useful_images.append(image_path)
                        self.logger.info(f"✔ Relevant Page: {image_path}")

                except json.JSONDecodeError:
                    self.logger.error(f"Invalid JSON format in response: {answer}")

            except requests.exceptions.RequestException as e:
                self.logger.error(f"VLM API request failed: {e}")

        # If relevant pages exist, save them as a new processed PDF in `processed_attachments`
        if useful_images:
            processed_pdf_path = os.path.join(processed_attachments_path, os.path.basename(attachment_path).replace(".pdf", "_processed.pdf"))
            self.combine_images_into_pdf(useful_images, processed_pdf_path)
            self.logger.info(f"Processed PDF saved at: {processed_pdf_path}")

            # Delete temporary PDF if it was created
            if temp_pdf_path and os.path.exists(temp_pdf_path):
                os.remove(temp_pdf_path)
                self.logger.info(f"Deleted temporary PDF: {temp_pdf_path}")

            return processed_pdf_path

        # If no relevant pages were found, delete the converted PDF
        if temp_pdf_path and os.path.exists(temp_pdf_path):
            shutil.rmtree(temp_pdf_path, ignore_errors=True)
            self.logger.info(f"Deleted temporary PDF: {temp_pdf_path}")

        self.logger.info(f"No relevant content found in {attachment_path}")
        return "not_relevant"

    def combine_images_into_pdf(self, useful_images: List[str], attachment_path: str) -> str:
        """
        Combines relevant pages (images) into a single PDF.

        Args:
            useful_images (List[str]): List of image paths for relevant pages.
            attachment_path (str): The original attachment file path.

        Returns:
            str: Path to the final processed PDF.
        """
        if not useful_images:
            self.logger.info(
                f"No relevant pages found for {attachment_path}. Skipping PDF creation."
            )
            return ""

        processed_folder = os.path.join(
            os.path.dirname(os.path.dirname(attachment_path)), "processed_attachments"
        )
        os.makedirs(processed_folder, exist_ok=True)
        output_pdf_path = os.path.join(
            processed_folder,
            os.path.basename(attachment_path),
        )

        pdf = FPDF()
        for image_path in useful_images:
            pdf.add_page()
            pdf.image(image_path, x=0, y=0, w=210, h=297)

        pdf.output(output_pdf_path, "F")
        self.logger.info(f"Saved processed PDF: {output_pdf_path}")

        temp_image_dir = os.path.dirname(useful_images[0])
        shutil.rmtree(temp_image_dir, ignore_errors=True)

        return output_pdf_path

    def extract_images_from_docx(
        docx_path: str, output_folder: str = None
    ) -> List[str]:
        """
        Extracts images from a DOCX file and saves them as image files.

        Args:
            docx_path (str): Path to the DOCX file.
            output_folder (str, optional): Folder to save extracted images.

        Returns:
            List[str]: List of paths to the saved image files.
        """
        doc = docx.Document(docx_path)

        if output_folder is None:
            output_folder = os.path.splitext(docx_path)[0] + "_images"
        os.makedirs(output_folder, exist_ok=True)

        image_paths = []
        image_count = 0

        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].target_ref:
                image_data = doc.part.rels[rel].target_part.blob

                image = Image.open(io.BytesIO(image_data))
                image_filename = f"docx_image_{image_count+1}.png"
                image_path = os.path.join(output_folder, image_filename)
                image.save(image_path, "PNG")

                image_paths.append(image_path)
                image_count += 1

        return image_paths

    def extract_text_from_pdf(self, file_path: str):
        """Extracts text from a PDF file."""
        """Extracts text from a scanned PDF using OCR."""
        extracted_text = ""

        # Convert PDF pages to images
        images = convert_from_path(file_path, poppler_path="/opt/homebrew/bin")  # Adjust poppler_path if needed

        for i, image in enumerate(images):
            text = pytesseract.image_to_string(image)
            extracted_text += f"\nPage {i+1}\n{text.strip()}\n\n"
            print(f"Extracted OCR Text from Page {i+1}:\n{text}")
        print(extracted_text)
        return extracted_text

    def extract_images_from_pdf(self, file_path: str, min_contour_area=5000):
        """
        Extracts images from a PDF file.

        Args:
            file_path (str): Path to the PDF file.
            min_contour_area (int): Minimum contour area to consider an object as a figure.

        Returns:
            dict: A dictionary containing extracted figures.
        """
        save_directory = os.path.join(os.getcwd(), "extracted_images")
        if os.path.exists(save_directory):
            for file in os.listdir(save_directory):
                os.remove(os.path.join(save_directory, file))  # Clear the directory before saving new images
        os.makedirs(save_directory, exist_ok=True)

        poppler_path = "/opt/homebrew/bin"
        images = convert_from_path(file_path, poppler_path=poppler_path)
        extracted_figures = {}
        all_extracted_images = []

        for page_num, image in enumerate(images):
            try:
                if image is None:
                    self.logger.warning(f"Skipping page {page_num+1} in {file_path} (No image content).")
                    continue

                open_cv_image = np.array(image.convert("RGB"))
                print(f"Processing Page {page_num+1} of {file_path}")

                if open_cv_image is None or open_cv_image.size == 0:
                    self.logger.warning(f"Skipping page {page_num+1} in {file_path} (Invalid image data).")
                    continue

                if len(open_cv_image.shape) == 3:
                    open_cv_image = cv2.cvtColor(open_cv_image, cv2.COLOR_RGB2BGR)
                else:
                    self.logger.warning(f"Skipping page {page_num+1}: Unexpected image shape {open_cv_image.shape}")
                    continue

                gray = cv2.cvtColor(open_cv_image, cv2.COLOR_BGR2GRAY)
                edges = cv2.Canny(gray, 50, 150)

                contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
                figure_paths = []

                for idx, contour in enumerate(contours):
                    if cv2.contourArea(contour) < min_contour_area:
                        continue
                    x, y, w, h = cv2.boundingRect(contour)
                    figure_image = image.crop((x, y, x + w, y + h))
                    figure_path = os.path.join(save_directory, f"page_{page_num+1}_figure_{idx+1}.png")
                    figure_image.save(figure_path, "PNG")
                    figure_paths.append(figure_path)

                all_extracted_images.extend(figure_paths)

            except Exception as e:
                self.logger.error(f"Error processing page {page_num+1} of {file_path}: {e}")

        # Filter images after extraction
        useful_figures = vlm.filter_images(all_extracted_images)

        # Remove non-useful images and retain only useful ones
        for file_path in all_extracted_images:
            if file_path not in useful_figures:
                os.remove(file_path)

        extracted_figures = {os.path.basename(img_path).split("_figure_")[0].replace("page_", "Page "): [] for img_path in useful_figures}
        for img_path in useful_figures:
            page_info = os.path.basename(img_path).split("_figure_")[0].replace("page_", "Page ")
            extracted_figures[page_info].append(img_path)

        return extracted_figures

    def extract_text_from_docx(self, file_path: str):
        """Extracts text from a Word document."""
        doc = Document(file_path)
        full_text = []

        for para in doc.paragraphs:
            full_text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                full_text.append("\t".join(row_text))
        print('\n'.join(full_text))
        return '\n'.join(full_text)
  
    def separate_text_and_images(self, file_path: str):
        """
        Extracts both text and images from an attachment.
        Returns a tuple: (extracted_text, list_of_images).
        """
        mime_type, _ = mimetypes.guess_type(file_path)
        text = ""
        images = []

        if mime_type and "pdf" in mime_type:
            text = self.extract_text_from_pdf(file_path)
            # images = self.extract_images_from_pdf(file_path)
            images = self.extract_images_from_pdf_without_filtering(file_path)
        elif mime_type and "word" in mime_type or file_path.endswith(".docx"):
            text = self.extract_text_from_docx(file_path)
            images = self.extract_images_from_docx(file_path)
        return text, images, file_path
    
    def chunk_text(self, text):
        os.environ["TOKENIZERS_PARALLELISM"] = "false"
        semantic_chunker = SemanticChunker(
            embedding_model, breakpoint_threshold_type="percentile"
        )
        semantic_chunks = semantic_chunker.create_documents([text])
        print(len(semantic_chunks))

        print([chunk.page_content for chunk in semantic_chunks])
        return [chunk.page_content for chunk in semantic_chunks]
    
    def process_and_save_attachments(self):
        """
        Processes filtered useful attachments:
        - Extracts text and images.
        - Generates semantic chunks from text.
        - Saves text chunks and images as pickle files in 'data_pkl'.
        """
        print("filtering useful attachments")
        useful_attachments = self.filter_useful_attachments()
        print("filtered useful attachments")
        if not useful_attachments:
            self.logger.info("No useful attachments found.")
            return

        # Define root directory for pickle storage
        root_pkl_dir = os.path.join(os.getcwd(), "data_pkl")
        os.makedirs(root_pkl_dir, exist_ok=True)

        for attachment_path in useful_attachments:
            attachment_name = os.path.splitext(os.path.basename(attachment_path))[0]
            save_dir = os.path.join(root_pkl_dir, attachment_name)
            os.makedirs(save_dir, exist_ok=True)

            # Extract text and images
            print("extracting text and images")
            extracted_text, extracted_images, _ = self.separate_text_and_images(attachment_path)

            # Generate semantic text chunks
            text_chunks = self.chunk_text(extracted_text)

            # Save text chunks as pickle
            text_pickle_path = os.path.join(save_dir, "text_chunks.pkl")
            with open(text_pickle_path, "wb") as f:
                pickle.dump(text_chunks, f)

            # Save images as pickle
            images_pickle_path = os.path.join(save_dir, "images.pkl")
            with open(images_pickle_path, "wb") as f:
                pickle.dump(extracted_images, f)

            self.logger.info(f"Processed and saved: {attachment_name}")
            self.logger.info(f"  - Text chunks: {text_pickle_path}")
            self.logger.info(f"  - Images: {images_pickle_path}")
        print("Finished processing all attachments.")

    # def process_and_save_attachments(self):
    #     """
    #     Accesses processed attachments stored in each email's `processed_attachments/` directory,
    #     extracts text and images, generates semantic chunks, and saves them in `data_pkl/`.
        
    #     - Text is printed to the console.
    #     - Text chunks are saved as Pickle.
    #     - Images are saved as PNG files inside `images/` subdirectory.
    #     """
    #     print("Accessing processed attachments...")
        
    #     if not os.path.exists(self.directory):
    #         self.logger.warning(f"Directory {self.directory} does not exist.")
    #         return

    #     # Define root directory for storage
    #     root_pkl_dir = os.path.join(os.getcwd(), "data_pkl")
    #     os.makedirs(root_pkl_dir, exist_ok=True)

    #     for email_folder in os.listdir(self.directory):
    #         email_path = os.path.join(self.directory, email_folder)
    #         processed_attachments_path = os.path.join(email_path, "processed_attachments")

    #         if not os.path.isdir(email_path):
    #             continue

    #         if not os.path.exists(processed_attachments_path):
    #             self.logger.warning(f"No processed attachments found in: {processed_attachments_path}")
    #             continue

    #         for file in os.listdir(processed_attachments_path):
    #             attachment_path = os.path.join(processed_attachments_path, file)
    #             attachment_name = os.path.splitext(file)[0]

    #             save_dir = os.path.join(root_pkl_dir, email_folder, attachment_name)
    #             images_dir = os.path.join(save_dir, "images")  # Save images in a subdirectory
    #             os.makedirs(images_dir, exist_ok=True)

    #             # Extract text and images
    #             print(f"Extracting text and images from: {attachment_path}")
    #             extracted_text, extracted_images, _ = self.separate_text_and_images(attachment_path)

    #             # Print extracted text
    #             print(f"\n📄 Extracted Text from {file}:\n")
    #             print(extracted_text)
    #             print("=" * 80)  # Separator for readability

    #             # Generate semantic text chunks
    #             text_chunks = self.chunk_text(extracted_text)

    #             # Save text chunks as pickle
    #             text_pickle_path = os.path.join(save_dir, "text_chunks.pkl")
    #             with open(text_pickle_path, "wb") as f:
    #                 pickle.dump(text_chunks, f)

    #             # Save images as PNG files
    #             for page_num, img_paths in extracted_images.items():
    #                 for img_idx, img_path in enumerate(img_paths):
    #                     img = Image.open(img_path)
    #                     img_save_path = os.path.join(images_dir, f"{page_num}_img_{img_idx+1}.png")
    #                     img.save(img_save_path, "PNG")

    #             self.logger.info(f"Processed and saved: {attachment_name}")
    #             self.logger.info(f"  - Text chunks: {text_pickle_path}")
    #             self.logger.info(f"  - Images saved in: {images_dir}")

    #     print("Finished processing all attachments.")

    def extract_images_from_pdf_without_filtering(self, file_path: str, min_contour_area=5000):
        """
        Extracts images from a PDF file.

        Args:
            file_path (str): Path to the PDF file.
            min_contour_area (int): Minimum contour area to consider an object as a figure.

        Returns:
            dict: A dictionary containing extracted figures.
        """
        save_directory = os.path.join(os.getcwd(), "extracted_images")
        if os.path.exists(save_directory):
            for file in os.listdir(save_directory):
                os.remove(os.path.join(save_directory, file))  # Clear the directory before saving new images
        os.makedirs(save_directory, exist_ok=True)

        poppler_path = "/opt/homebrew/bin"
        images = convert_from_path(file_path, poppler_path=poppler_path)
        extracted_figures = {}
        all_extracted_images = []

        for page_num, image in enumerate(images):
            try:
                if image is None:
                    self.logger.warning(f"Skipping page {page_num+1} in {file_path} (No image content).")
                    continue

                open_cv_image = np.array(image.convert("RGB"))
                print(f"Processing Page {page_num+1} of {file_path}")

                if open_cv_image is None or open_cv_image.size == 0:
                    self.logger.warning(f"Skipping page {page_num+1} in {file_path} (Invalid image data).")
                    continue

                if len(open_cv_image.shape) == 3:
                    open_cv_image = cv2.cvtColor(open_cv_image, cv2.COLOR_RGB2BGR)
                else:
                    self.logger.warning(f"Skipping page {page_num+1}: Unexpected image shape {open_cv_image.shape}")
                    continue

                gray = cv2.cvtColor(open_cv_image, cv2.COLOR_BGR2GRAY)
                edges = cv2.Canny(gray, 50, 150)

                contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
                figure_paths = []

                for idx, contour in enumerate(contours):
                    if cv2.contourArea(contour) < min_contour_area:
                        continue
                    x, y, w, h = cv2.boundingRect(contour)
                    figure_image = image.crop((x, y, x + w, y + h))
                    figure_path = os.path.join(save_directory, f"page_{page_num+1}_figure_{idx+1}.png")
                    figure_image.save(figure_path, "PNG")
                    figure_paths.append(figure_path)

                all_extracted_images.extend(figure_paths)

            except Exception as e:
                self.logger.error(f"Error processing page {page_num+1} of {file_path}: {e}")

        extracted_figures = {os.path.basename(img_path).split("_figure_")[0].replace("page_", "Page "): [] for img_path in all_extracted_images}
        for img_path in all_extracted_images:
            page_info = os.path.basename(img_path).split("_figure_")[0].replace("page_", "Page ")
            extracted_figures[page_info].append(img_path)

        return extracted_figures

    def process_user_uploads(self, file_path: str):
        """
        Processes a user-uploaded document:
        1. Extracts text and images using the `separate_text_and_images` method.
        2. Chunks the extracted text using `chunk_text`.
        3. Generates image summaries using `vlm.generate_image_summaries()`.
        4. Returns the list of text chunks and image summaries.
        
        Args:
            file_path (str): The path to the uploaded file.
        
        Returns:
            Tuple[List[str], List[str]]: A tuple containing a list of text chunks and a list of image summaries.
        """
        document_parser = DocumentParser()
        
        extracted_text, extracted_images, _ = document_parser.separate_text_and_images(file_path)
        
        text_chunks = document_parser.chunk_text(extracted_text)

        # image_summaries = vlm.generate_image_summaries(extracted_images)

        return text_chunks
